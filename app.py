import os
import sys
import logging
from datetime import datetime
from dotenv import load_dotenv
from flask import Flask, jsonify, send_from_directory, request
from flask_cors import CORS
from flask_apscheduler import APScheduler
import chinese_calendar
import json
import threading

from data_fetcher import DataFetcher
from ai_analyzer import AIAnalyzer
from notifier import DingTalkNotifier

# 加载环境
load_dotenv()

# 日志配置
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("app.log", encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("App")

class Config:
    SCHEDULER_API_ENABLED = True

app = Flask(__name__)
CORS(app) # 启用跨域支持
app.config.from_object(Config())

scheduler = APScheduler()
scheduler.init_app(app)
scheduler.start()

fetcher = DataFetcher()
analyzer = AIAnalyzer()
notifier = DingTalkNotifier()

# 互斥锁，防止并发执行任务
task_lock = threading.Lock()
current_task_status = {"running": False, "last_result": None}

def daily_job(manual=False, send_notify=True):
    """每日执行的核心任务"""
    if not task_lock.acquire(blocking=False):
        logger.warning("任务已在运行中，跳过本次请求。")
        return False
        
    try:
        current_task_status["running"] = True
        now = datetime.now()
        
        # 如果是自动触发，检查节假日
        if not manual and chinese_calendar.is_holiday(now):
            logger.info(f"今日 {now.strftime('%Y-%m-%d')} 是节假日，跳过自动推送。")
            return True

        logger.info(f"{'手动' if manual else '自动'}触发研报生成任务 (推送: {send_notify})...")
        
        # 1. 采集数据
        target_symbols = [s.strip() for s in os.getenv("TARGET_SYMBOLS", "").split(",") if s.strip()]
        quotes_data = fetcher.get_futures_quotes(target_symbols)
        news_df = fetcher.get_futures_news()
        
        news_list = []
        if news_df is not None:
            for idx, row in news_df.iterrows():
                content = row.get('title', row.get('content', '无内容'))
                time_str = row.get('pubDate', row.get('发布时间', '无时间'))
                news_list.append(f"{time_str}：{content}")
        news_str = "\n".join(news_list)

        # 2. AI 分析
        report = analyzer.generate_report(quotes_data, news_str)
        
        # 3. 持久化到状态文件 (供前端显示)
        state_file = 'workflow_state.json'
        state = {}
        if os.path.exists(state_file):
            try:
                with open(state_file, 'r', encoding='utf-8') as f:
                    state = json.load(f)
            except:
                pass
        
        # 更新研报列表 (保留最近 5 条)
        reports = state.get("reports", [])
        new_report = {
            "title": f"研报 {now.strftime('%Y-%m-%d %H:%M')}",
            "content": report,
            "date": now.strftime('%Y-%m-%d %H:%M:%S')
        }
        reports.insert(0, new_report)
        state["reports"] = reports[:5]
        state["quotes_data"] = quotes_data # 同步更新行情数据
        state["date_str"] = now.strftime('%m-%d')
        
        with open(state_file, 'w', encoding='utf-8') as f:
            json.dump(state, f, ensure_ascii=False, indent=2)

        # 4. 推送 (仅当需要时)
        if send_notify:
            title = f"【期货研报-每日智领】({now.strftime('%m-%d %H:%M')})"
            full_text = f"# {title}\n\n{report}"
            notifier.send_markdown(title, full_text)
            logger.info("钉钉推送已发送。")
        else:
            logger.info("已跳过钉钉推送。")
        
        current_task_status["last_result"] = "Success"
        logger.info("任务执行完毕。")
        return True
    except Exception as e:
        logger.exception("任务执行出错")
        current_task_status["last_result"] = f"Error: {str(e)}"
        return False
    finally:
        current_task_status["running"] = False
        task_lock.release()


@app.route('/api/state')
def get_state():
    """获取当前系统状态（从 workflow_state.json 读取）"""
    try:
        with open('workflow_state.json', 'r', encoding='utf-8') as f:
            state = json.load(f)
        return jsonify({
            "status": "Success",
            "data": state,
            "system": current_task_status
        })
    except Exception as e:
        return jsonify({"status": "Error", "message": str(e)}), 500

@app.route('/api/trigger', methods=['POST'])
def trigger_task():
    """手动触发任务接口"""
    if current_task_status["running"]:
        return jsonify({"status": "Busy", "message": "任务正在执行中"}), 409
    
    # 异步执行
    thread = threading.Thread(target=daily_job, kwargs={"manual": True, "send_notify": False})

    thread.start()
    return jsonify({"status": "Started", "message": "任务已在后台启动"})

@app.route('/api/logs')
def get_logs():
    """获取最新的日志内容"""
    try:
        if os.path.exists("app.log"):
            with open("app.log", "r", encoding='utf-8') as f:
                # 只取最后 100 行
                lines = f.readlines()
                return jsonify({"status": "Success", "logs": lines[-100:]})
        return jsonify({"status": "Success", "logs": []})
    except Exception as e:
        return jsonify({"status": "Error", "message": str(e)}), 500

@app.route('/health')
def health():
    return jsonify({"status": "OK", "time": datetime.now().isoformat()})

@app.route('/api/news')
def get_all_news():
    """获取综合财经新闻与国际局势"""
    try:
        # 定义核心板块关键词（涵盖所有 16 个品种及宏观词）
        commodity_keywords = [
            "豆粕", "玉米", "菜粕", "生猪", "白糖", "棉花", 
            "甲醇", "尿素", "PTA", "纯碱", "PVC", "淀粉", 
            "大豆", "豆油", "乙醇", "生物柴",
            "CBOT", "USDA", "美联储", "地缘", "冲突"
        ]
        news_list = fetcher.get_commodity_news(commodity_keywords)
        return jsonify({"status": "Success", "data": news_list})
    except Exception as e:
        logger.error(f"获取综合新闻失败: {e}")
        return jsonify({"status": "Error", "message": str(e)}), 500

@app.route('/api/analyze_news', methods=['POST'])
def analyze_current_news():
    """使用 AI 分析当前新闻"""
    try:
        data = request.json
        news_items = data.get('news', [])
        if not news_items:
            return jsonify({"status": "Error", "message": "No news data provided"}), 400
        
        analysis = analyzer.analyze_news_impact(news_items)
        return jsonify({"status": "Success", "analysis": analysis})
    except Exception as e:
        logger.error(f"AI 新闻分析失败: {e}")
        return jsonify({"status": "Error", "message": str(e)}), 500

@app.route('/api/history')
def get_history():
    """获取品种历史走势数据"""
    symbol = request.args.get('symbol', '豆粕')
    try:
        df = fetcher.get_futures_history(symbol, days=30)
        if df is not None:
            # 格式化为前端 Recharts 易读的 JSON
            history_data = []
            for _, row in df.iterrows():
                history_data.append({
                    "name": str(row['date']).split()[0][-5:], # 取月-日
                    "value": float(row['close'])
                })
            return jsonify({"status": "Success", "data": history_data})
        return jsonify({"status": "Error", "message": "未找到历史数据"}), 404
    except Exception as e:
        return jsonify({"status": "Error", "message": str(e)}), 500

@app.route('/api/market_analysis')
def get_market_analysis():
    """获取综合行情分析数据"""
    try:
        # 获取核心品种
        target_symbols = ["豆粕", "玉米", "大豆", "白糖", "棉花", "甲醇", "PTA"]
        analysis_data = fetcher.get_market_analysis_data(target_symbols)
        if analysis_data:
            return jsonify({"status": "Success", "data": analysis_data})
        return jsonify({"status": "Error", "message": "计算分析数据失败"}), 500
    except Exception as e:
        logger.error(f"行情分析接口异常: {e}")
        return jsonify({"status": "Error", "message": str(e)}), 500

@app.route('/api/inventory')
def get_inventory():
    """获取品种库存及其季节性数据"""
    symbol = request.args.get('symbol', '豆粕')
    try:
        data = fetcher.get_inventory_seasonality(symbol)
        if data:
            return jsonify({"status": "Success", "data": data})
        
        # 降级：如果无季节性数据，尝试获取基础数据
        basic_data = fetcher.get_futures_inventory(symbol)
        if basic_data:
             return jsonify({
                 "status": "Success", 
                 "data": {
                     "current": basic_data, 
                     "history": [], 
                     "symbol": symbol,
                     "unit": "手"
                 }
             })
             
        return jsonify({"status": "Error", "message": "未找到库存数据"}), 404
    except Exception as e:
        return jsonify({"status": "Error", "message": str(e)}), 500


@app.route('/api/agricultural_quotes')
def get_agricultural_quotes():
    """获取所有农产品期货行情数据"""
    try:
        # 获取所有农产品期货品种
        agri_symbols = fetcher.agricultural_symbols
        quotes_data = fetcher.get_futures_quotes(agri_symbols)
        
        # 添加分类信息
        for item in quotes_data:
            item['category'] = fetcher.category_map.get(item['symbol'], '其他')
        
        return jsonify({
            "status": "Success", 
            "data": quotes_data,
            "categories": list(set(fetcher.category_map.values()))
        })
    except Exception as e:
        logger.error(f"获取农产品行情失败: {e}")
        return jsonify({"status": "Error", "message": str(e)}), 500


@app.route('/api/all_symbols')
def get_all_symbols():
    """获取所有支持的期货品种列表"""
    try:
        return jsonify({
            "status": "Success",
            "data": {
                "agricultural": fetcher.agricultural_symbols,
                "all": list(fetcher.symbol_map.keys()),
                "categories": {k: v for k, v in fetcher.category_map.items()}
            }
        })
    except Exception as e:
        logger.error(f"获取品种列表失败: {e}")
        return jsonify({"status": "Error", "message": str(e)}), 500


@app.route('/api/technical_indicators')
def get_technical_indicators():
    """获取品种技术指标数据"""
    symbol = request.args.get('symbol', '豆粕')
    days = int(request.args.get('days', 60))
    try:
        data = fetcher.calculate_technical_indicators(symbol, days)
        if data:
            return jsonify({"status": "Success", "data": data})
        return jsonify({"status": "Error", "message": "无法获取技术指标数据"}), 404
    except Exception as e:
        logger.error(f"技术指标接口异常: {e}")
        return jsonify({"status": "Error", "message": str(e)}), 500


@app.route('/api/spread_analysis')
def get_spread_analysis():
    """获取两个品种的价差分析"""
    symbol1 = request.args.get('symbol1', '豆粕')
    symbol2 = request.args.get('symbol2', '玉米')
    days = int(request.args.get('days', 30))
    try:
        data = fetcher.get_spread_analysis(symbol1, symbol2, days)
        if data:
            return jsonify({"status": "Success", "data": data})
        return jsonify({"status": "Error", "message": "无法获取价差数据"}), 404
    except Exception as e:
        logger.error(f"价差分析接口异常: {e}")
        return jsonify({"status": "Error", "message": str(e)}), 500


@app.route('/api/price_ratio')
def get_price_ratio():
    """获取两个品种的比价走势"""
    symbol1 = request.args.get('symbol1', '豆粕')
    symbol2 = request.args.get('symbol2', '玉米')
    days = int(request.args.get('days', 30))
    try:
        data = fetcher.get_price_ratio(symbol1, symbol2, days)
        if data:
            return jsonify({"status": "Success", "data": data})
        return jsonify({"status": "Error", "message": "无法获取比价数据"}), 404
    except Exception as e:
        logger.error(f"比价分析接口异常: {e}")
        return jsonify({"status": "Error", "message": str(e)}), 500


@app.route('/api/favorites', methods=['GET', 'POST', 'DELETE'])
def manage_favorites():
    """管理自选品种"""
    fav_file = 'favorites.json'
    
    try:
        # 读取现有数据
        favorites = []
        if os.path.exists(fav_file):
            with open(fav_file, 'r', encoding='utf-8') as f:
                favorites = json.load(f)
        
        if request.method == 'GET':
            return jsonify({"status": "Success", "data": favorites})
        
        elif request.method == 'POST':
            data = request.json
            symbol = data.get('symbol')
            if not symbol:
                return jsonify({"status": "Error", "message": "缺少品种参数"}), 400
            
            if symbol not in favorites:
                favorites.append(symbol)
                with open(fav_file, 'w', encoding='utf-8') as f:
                    json.dump(favorites, f, ensure_ascii=False)
            
            return jsonify({"status": "Success", "data": favorites})
        
        elif request.method == 'DELETE':
            data = request.json
            symbol = data.get('symbol')
            if symbol in favorites:
                favorites.remove(symbol)
                with open(fav_file, 'w', encoding='utf-8') as f:
                    json.dump(favorites, f, ensure_ascii=False)
            
            return jsonify({"status": "Success", "data": favorites})
    
    except Exception as e:
        logger.error(f"自选管理异常: {e}")
        return jsonify({"status": "Error", "message": str(e)}), 500


@app.route('/api/price_alerts', methods=['GET', 'POST', 'DELETE'])
def manage_price_alerts():
    """管理价格预警"""
    alert_file = 'price_alerts.json'
    
    try:
        # 读取现有数据
        alerts = []
        if os.path.exists(alert_file):
            with open(alert_file, 'r', encoding='utf-8') as f:
                alerts = json.load(f)
        
        if request.method == 'GET':
            return jsonify({"status": "Success", "data": alerts})
        
        elif request.method == 'POST':
            data = request.json
            symbol = data.get('symbol')
            alert_type = data.get('type')  # 'above' or 'below'
            price = data.get('price')
            
            if not symbol or not alert_type or not price:
                return jsonify({"status": "Error", "message": "参数不完整"}), 400
            
            new_alert = {
                "id": len(alerts) + 1,
                "symbol": symbol,
                "type": alert_type,
                "price": float(price),
                "enabled": True,
                "created_at": datetime.now().strftime('%Y-%m-%d %H:%M')
            }
            
            alerts.append(new_alert)
            with open(alert_file, 'w', encoding='utf-8') as f:
                json.dump(alerts, f, ensure_ascii=False, indent=2)
            
            return jsonify({"status": "Success", "data": alerts})
        
        elif request.method == 'DELETE':
            alert_id = request.args.get('id')
            alerts = [a for a in alerts if a.get('id') != int(alert_id) if alert_id]
            with open(alert_file, 'w', encoding='utf-8') as f:
                json.dump(alerts, f, ensure_ascii=False, indent=2)
            
            return jsonify({"status": "Success", "data": alerts})
    
    except Exception as e:
        logger.error(f"预警管理异常: {e}")
        return jsonify({"status": "Error", "message": str(e)}), 500


@app.route('/api/check_alerts')
def check_price_alerts():
    """检查价格预警是否触发"""
    try:
        alert_file = 'price_alerts.json'
        if not os.path.exists(alert_file):
            return jsonify({"status": "Success", "data": [], "triggered": []})
        
        with open(alert_file, 'r', encoding='utf-8') as f:
            alerts = json.load(f)
        
        # 获取当前价格
        symbols = list(set(a['symbol'] for a in alerts))
        quotes = fetcher.get_futures_quotes(symbols)
        current_prices = {q['symbol']: q['price'] for q in quotes}
        
        triggered = []
        for alert in alerts:
            if not alert.get('enabled', True):
                continue
            symbol = alert['symbol']
            if symbol not in current_prices:
                continue
            
            current = current_prices[symbol]
            if alert['type'] == 'above' and current >= alert['price']:
                triggered.append({**alert, "current_price": current})
            elif alert['type'] == 'below' and current <= alert['price']:
                triggered.append({**alert, "current_price": current})
        
        return jsonify({
            "status": "Success", 
            "data": alerts,
            "triggered": triggered,
            "current_prices": current_prices
        })
    except Exception as e:
        logger.error(f"预警检查异常: {e}")
        return jsonify({"status": "Error", "message": str(e)}), 500


@app.route('/static/charts/<path:filename>')
def serve_charts(filename):
    return send_from_directory('static/charts', filename)


# ============= 研报管理 API =============
REPORTS_FILE = 'reports_data.json'

def load_reports():
    """加载研报数据"""
    if not os.path.exists(REPORTS_FILE):
        return []
    try:
        with open(REPORTS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except:
        return []

def save_reports(reports):
    """保存研报数据"""
    with open(REPORTS_FILE, 'w', encoding='utf-8') as f:
        json.dump(reports, f, ensure_ascii=False, indent=2)

@app.route('/api/reports', methods=['GET'])
def get_reports():
    """获取所有研报列表"""
    try:
        reports = load_reports()
        # 按发布日期倒序排列
        reports.sort(key=lambda x: x.get('publishDate', ''), reverse=True)
        return jsonify({"status": "Success", "data": reports})
    except Exception as e:
        logger.error(f"获取研报列表异常：{e}")
        return jsonify({"status": "Error", "message": str(e)}), 500

@app.route('/api/reports', methods=['POST'])
def create_report():
    """创建新研报（支持上传 PDF/Word 文件）"""
    try:
        # 获取表单数据
        title = request.form.get('title')
        summary = request.form.get('summary', '')
        category = request.form.get('category', '行业')
        author = request.form.get('author', '用户上传')
        risk_level = request.form.get('risk_level', '中')
        target_symbols = request.form.get('target_symbols', '')
        
        if not title:
            return jsonify({"status": "Error", "message": "标题不能为空"}), 400
        
        # 处理文件上传
        file_path = None
        file_type = None
        if 'file' in request.files:
            file = request.files['file']
            if file and file.filename:
                # 验证文件类型
                allowed_extensions = {'pdf', 'doc', 'docx'}
                ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
                
                if ext not in allowed_extensions:
                    return jsonify({"status": "Error", "message": "仅支持 PDF 和 Word 文件"}), 400
                
                # 创建上传目录
                upload_dir = 'uploads/reports'
                os.makedirs(upload_dir, exist_ok=True)
                
                # 生成唯一文件名
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"report_{timestamp}_{file.filename}"
                file_path = os.path.join(upload_dir, filename)
                file.save(file_path)
                file_type = ext
        
        # 解析目标品种
        symbols_list = [s.strip() for s in target_symbols.split(',') if s.strip()]
        
        # 创建研报对象
        new_report = {
            "id": f"report_{datetime.now().strftime('%Y%m%d%H%M%S')}",
            "title": title,
            "summary": summary,
            "content": request.form.get('content', ''),
            "category": category,
            "author": author,
            "publishDate": datetime.now().strftime('%Y-%m-%d'),
            "readTime": len(request.form.get('content', '')) // 500 + 1,
            "views": 0,
            "rating": 0,
            "tags": [category, '用户上传'],
            "riskLevel": risk_level,
            "targetSymbols": symbols_list if symbols_list else ['所有品种'],
            "fileType": file_type,
            "filePath": file_path,
            "createdAt": datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        
        # 加载现有研报并添加新研报
        reports = load_reports()
        reports.insert(0, new_report)
        save_reports(reports)
        
        return jsonify({"status": "Success", "data": new_report})
    
    except Exception as e:
        logger.error(f"创建研报异常：{e}")
        return jsonify({"status": "Error", "message": str(e)}), 500

@app.route('/api/reports/<report_id>', methods=['PUT'])
def update_report(report_id):
    """更新研报"""
    try:
        reports = load_reports()
        report_index = next((i for i, r in enumerate(reports) if r['id'] == report_id), None)
        
        if report_index is None:
            return jsonify({"status": "Error", "message": "研报不存在"}), 404
        
        # 更新字段
        data = request.json
        report = reports[report_index]
        
        if 'title' in data:
            report['title'] = data['title']
        if 'summary' in data:
            report['summary'] = data['summary']
        if 'content' in data:
            report['content'] = data['content']
        if 'category' in data:
            report['category'] = data['category']
        if 'author' in data:
            report['author'] = data['author']
        if 'riskLevel' in data:
            report['riskLevel'] = data['riskLevel']
        if 'targetSymbols' in data:
            report['targetSymbols'] = data['targetSymbols']
        if 'tags' in data:
            report['tags'] = data['tags']
        
        report['updatedAt'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        reports[report_index] = report
        save_reports(reports)
        
        return jsonify({"status": "Success", "data": report})
    
    except Exception as e:
        logger.error(f"更新研报异常：{e}")
        return jsonify({"status": "Error", "message": str(e)}), 500

@app.route('/api/reports/<report_id>', methods=['DELETE'])
def delete_report(report_id):
    """删除研报"""
    try:
        reports = load_reports()
        report = next((r for r in reports if r['id'] == report_id), None)
        
        if report is None:
            return jsonify({"status": "Error", "message": "研报不存在"}), 404
        
        # 删除关联文件
        if report.get('filePath') and os.path.exists(report['filePath']):
            os.remove(report['filePath'])
        
        # 从列表中移除
        reports = [r for r in reports if r['id'] != report_id]
        save_reports(reports)
        
        return jsonify({"status": "Success", "message": "删除成功"})
    
    except Exception as e:
        logger.error(f"删除研报异常：{e}")
        return jsonify({"status": "Error", "message": str(e)}), 500

@app.route('/api/reports/<report_id>/view', methods=['POST'])
def increment_report_view(report_id):
    """增加研报阅读量"""
    try:
        reports = load_reports()
        report = next((r for r in reports if r['id'] == report_id), None)
        
        if report:
            report['views'] = report.get('views', 0) + 1
            save_reports(reports)
            return jsonify({"status": "Success", "data": report})
        
        return jsonify({"status": "Error", "message": "研报不存在"}), 404
    
    except Exception as e:
        logger.error(f"增加阅读量异常：{e}")
        return jsonify({"status": "Error", "message": str(e)}), 500

@app.route('/uploads/reports/<filename>')
def serve_report_file(filename):
    """提供研报文件下载"""
    return send_from_directory('uploads/reports', filename)

@app.route('/api/reports/<report_id>/preview', methods=['GET'])
def preview_report_file(report_id):
    """预览研报文件（PDF/Word 转 HTML）"""
    try:
        reports = load_reports()
        report = next((r for r in reports if r['id'] == report_id), None)
        
        if not report:
            return jsonify({"status": "Error", "message": "研报不存在"}), 404
        
        file_path = report.get('filePath')
        file_type = report.get('fileType')
        
        if not file_path or not os.path.exists(file_path):
            return jsonify({"status": "Error", "message": "文件不存在"}), 404
        
        # 根据文件类型进行预览
        if file_type == 'pdf':
            # 使用 PyMuPDF 将 PDF 转换为 HTML
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            total_pages = doc.page_count
            html_content = []
            
            # 转换每一页
            for page_num in range(min(total_pages, 5)):  # 只预览前 5 页
                page = doc[page_num]
                # 获取页面文本
                text = page.get_text()
                html_content.append(f"<div class='pdf-page' style='margin-bottom: 20px; padding: 20px; border: 1px solid #262630; background: #0a0a0a;'>")
                html_content.append(f"<div style='color: #5a5a5a; font-size: 12px; margin-bottom: 10px;'>第 {page_num + 1} 页</div>")
                html_content.append(f"<div style='color: #e8e6e3; white-space: pre-wrap; font-family: monospace;'>{text}</div>")
                html_content.append("</div>")
            
            doc.close()
            
            return jsonify({
                "status": "Success",
                "data": {
                    "type": "pdf",
                    "content": "".join(html_content),
                    "total_pages": total_pages,
                    "preview_pages": min(total_pages, 5)
                }
            })
        
        elif file_type in ['doc', 'docx']:
            # 使用 python-docx 读取 Word 文件
            from docx import Document
            
            doc = Document(file_path)
            html_content = []
            
            # 添加标题
            if doc.core_properties.title:
                html_content.append(f"<h2 style='color: #e8e6e3; margin-bottom: 20px;'>{doc.core_properties.title}</h2>")
            
            # 添加段落
            for i, para in enumerate(doc.paragraphs[:50]):  # 只预览前 50 段
                if para.text.strip():
                    html_content.append(f"<p style='color: #e8e6e3; margin: 10px 0; line-height: 1.6;'>{para.text}</p>")
            
            # 添加表格
            for table in doc.tables[:3]:  # 只预览前 3 个表格
                html_content.append("<table style='width: 100%; border-collapse: collapse; margin: 15px 0;'>")
                for row in table.rows:
                    html_content.append("<tr>")
                    for cell in row.cells:
                        html_content.append(f"<td style='border: 1px solid #262630; padding: 8px; color: #e8e6e3;'>{cell.text}</td>")
                    html_content.append("</tr>")
                html_content.append("</table>")
            
            return jsonify({
                "status": "Success",
                "data": {
                    "type": "word",
                    "content": "".join(html_content),
                    "paragraph_count": len(doc.paragraphs),
                    "preview_paragraphs": min(len(doc.paragraphs), 50)
                }
            })
        
        else:
            return jsonify({"status": "Error", "message": "不支持的文件类型"}), 400
    
    except Exception as e:
        logger.error(f"预览文件异常：{e}")
        return jsonify({"status": "Error", "message": str(e)}), 500

if __name__ == "__main__":
    # 添加定时计划：周一至周五 8:00（早晨）
    try:
        scheduler.add_job(id='daily_push', func=daily_job, trigger='cron', 
                          day_of_week='mon-fri', hour=8, minute=0,
                          kwargs={"send_notify": False})

    except Exception as e:
        logger.error(f"添加定时任务失败：{e}")
    
    logger.info("Flask 服务启动...")
    app.run(host='0.0.0.0', port=5000, debug=True)
